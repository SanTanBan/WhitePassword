# https://stackoverflow.com/questions/41979095/write-text-in-particular-font-color-in-ms-word-using-python-docx

from docx import Document
from docx.shared import RGBColor
import random
import matplotlib.pyplot as plt
import numpy as np
from docx2pdf import convert
import os
import pyperclip as pc


all_unicode = []
for i in range(256):
    all_unicode.append(chr(i))

alphabets=[]
for i in range(97,123):
    alphabets.append(chr(i))

ALPHABETS=[]
for i in range(65,91):
    ALPHABETS.append(chr(i))

numbers=[]
for i in range(48,58):
    numbers.append(chr(i))

special_characters=[' ','!','#','$','%','&',"`",'(',')','*','+',',','-','.','/',':',';','<','=','>','?','@','[',"\\","]",'^','_','{','|',"'",'}','~','"']



intent=input("\n Some reference of this password for your recollection. You may include name of the website: \t")
os.mkdir(intent)
intent+="/"
docu = Document()
paragraph = docu.add_paragraph()

print("\n It is preferable to randomize the length of the password you enter by a small margin for generating more uncertainity and preventing easy password cracks")
decision_of_randomizing_length=input("\n Would you like to randomize your password's length? Enter 0 ( ZERO ) if NO : \t")

if decision_of_randomizing_length=="0":
    length_of_password_requested=int(input("\n Exact length of the password: \t"))
else:
    length_of_password_entered=int(input("\n Approximate length of the password is: \t"))

    print("\n For Updated Security Resaons, the length of the password you just entered is to be considered as the Mean for generating a range within which the password length might lie...")
    print("\n As an example, if you entered 25, then it would be considered as a mean and the password range might be from 25-3 characters to 25+3 characters")
    print("\n The uncertainity range ( 3 in the above example ) would be used as a linear Deviation with your entered length as the average:")

    uncertainity_range_in_password_length=int(input("Enter Uncertainity Range: \t"))
    while uncertainity_range_in_password_length>=length_of_password_entered:
        print("The Uncertainity range must be less than the requested password length and is preferable to be a small value b/w 3 to 1")
        uncertainity_range_in_password_length=int(input("Enter Uncertainity Range: \t"))
    length_of_password_requested=random.randint(length_of_password_entered-uncertainity_range_in_password_length, length_of_password_entered+uncertainity_range_in_password_length)

# all_keyboard_entity={1: '`', 2: '1', 3: '2', 4: '3', 5: '4', 6: '5', 7: '6', 8: '7', 9: '8', 10: 'N', 11: '9', 12: '0', 13: '-', 14: '=', 15: 'q', 16: 'w', 17: 'e', 18: 'r', 19: 't', 20: 'y', 21: 'u', 22: 'i', 23: 'o', 24: 'p', 25: '[', 26: ']', 27: '\\', 28: 'a', 29: 's', 30: 'd', 31: 'f', 32: 'g', 33: 'h', 34: 'j', 35: 'k', 36: 'l', 37: ';', 38: "'", 39: 'z', 40: 'x', 41: 'c', 42: 'v', 43: 'b', 44: 'n', 45: 'm', 46: ',', 47: '.', 48: '/', 49: '~', 50: '!', 51: '@', 52: '#', 53: '$', 54: '%', 55: '^', 56: '&', 57: '*', 58: '', 59: '(', 60: ')', 61: '_', 62: '+', 63: '{', 64: '}', 65: '|', 66: ':', 67: '"', 68: '<', 69: '>', 70: '?', 71: 'Q', 72: 'W', 73: 'E', 74: 'R', 75: 'T', 76: 'Y', 77: 'U', 78: 'I', 79: 'O', 80: 'P', 81: 'A', 82: 'S', 83: 'D', 84: 'F', 85: 'G', 86: 'H', 87: 'J', 88: 'K', 89: 'L', 90: 'Z', 91: 'X', 92: 'C', 93: 'V', 94: 'B', 95: 'M', 96: ' '}

# No Space for one specific case # Comment the below line after using code once
all_keyboard_entity={1: '`', 2: '1', 3: '2', 4: '3', 5: '4', 6: '5', 7: '6', 8: '7', 9: '8', 10: 'N', 11: '9', 12: '0', 13: '-', 14: '=', 15: 'q', 16: 'w', 17: 'e', 18: 'r', 19: 't', 20: 'y', 21: 'u', 22: 'i', 23: 'o', 24: 'p', 25: '[', 26: ']', 27: '\\', 28: 'a', 29: 's', 30: 'd', 31: 'f', 32: 'g', 33: 'h', 34: 'j', 35: 'k', 36: 'l', 37: ';', 38: "'", 39: 'z', 40: 'x', 41: 'c', 42: 'v', 43: 'b', 44: 'n', 45: 'm', 46: ',', 47: '.', 48: '/', 49: '~', 50: '!', 51: '@', 52: '#', 53: '$', 54: '%', 55: '^', 56: '&', 57: '*', 58: '', 59: '(', 60: ')', 61: '_', 62: '+', 63: '{', 64: '}', 65: '|', 66: ':', 67: '"', 68: '<', 69: '>', 70: '?', 71: 'Q', 72: 'W', 73: 'E', 74: 'R', 75: 'T', 76: 'Y', 77: 'U', 78: 'I', 79: 'O', 80: 'P', 81: 'A', 82: 'S', 83: 'D', 84: 'F', 85: 'G', 86: 'H', 87: 'J', 88: 'K', 89: 'L', 90: 'Z', 91: 'X', 92: 'C', 93: 'V', 94: 'B', 95: 'M'}

num_of_options=len(all_keyboard_entity)
#print("Length of the KeyBoard Array is: ", num_of_options)

check=0
while check==0:

    PASSWORD=""
    check_number_in_password=0
    check_ALPHABET_in_password=0
    check_alphabet_in_password=0
    check_special_character_in_password=0

    for i in range(length_of_password_requested):
        rand=random.randint(1,num_of_options)
        PASSWORD+=all_keyboard_entity[rand]
        
        if all_keyboard_entity[rand] in alphabets:
            check_alphabet_in_password+=5
        elif all_keyboard_entity[rand] in ALPHABETS:
            check_ALPHABET_in_password+=5
        elif all_keyboard_entity[rand] in numbers:
            check_number_in_password+=5
        elif all_keyboard_entity[rand] in special_characters:
            check_special_character_in_password+=5

    if check_ALPHABET_in_password>0 and check_alphabet_in_password>0 and check_number_in_password>0 and check_number_in_password>0:
        check=1


run = paragraph.add_run('\n The Password is in within the arrows. Copy the everything within the arrows which is actually in white font. --=>>')
run.font.color.rgb = RGBColor(255, 0, 0)

run = paragraph.add_run(PASSWORD)
run.font.color.rgb = RGBColor(255,255,255)

run = paragraph.add_run('<<=-- Dont leave any space while copying everything within the arrows.')
run.font.color.rgb = RGBColor(255, 0, 0)

docu.save(intent+'White_Password.docx')
convert(intent+'White_Password.docx',intent+"White_Password_PDF.pdf")

print("\n The password files should not contain the information about where this password is being used, only the outside FOLDER NAME is to help the user regarding this.")

# After USING the password, clean CLIPBOARD
# If you dont wish to save the passwords then comment out the codes for saving the password in WORD and PDF files
# In this case you should be using the PASSWORD as an OTP through the CLIPBOARD
pc.copy(PASSWORD) #Copying the Password to the CLIPBOARD for immediate USE
pc.paste()

print("\n Your password is copied to your CLIPBOARD. USE EXTREME CAUTION !!!")