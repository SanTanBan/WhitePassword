"""
from docx import Document
from docx.shared import RGBColor
document = Document()
run = document.add_paragraph('some text').add_run()
font = run.font
font.color.rgb = RGBColor(0x42, 0x24, 0xE9)
p=document.add_paragraph('aaa')
document.save('demo1.docx')
"""
# https://stackoverflow.com/questions/41979095/write-text-in-particular-font-color-in-ms-word-using-python-docx

from docx import Document
from docx.shared import RGBColor
import random
import matplotlib.pyplot as plt
import numpy as np
from docx2pdf import convert


"""
run = paragraph.add_run('Red ')
run.font.color.rgb = RGBColor(255, 0, 0)
run = paragraph.add_run('Green ')
run.font.color.rgb = RGBColor(0x00, 0xFF, 0x00)
run = paragraph.add_run('Blue')
run.font.color.rgb = RGBColor.from_string('0000FF')
all_keyboard_entity={}
for i in range (1,95):
    all_keyboard_entity[i]=input("Enter Keyboard Element: ")
print(all_keyboard_entity)
"""
intent=input("Some reference of this password for your recollection. You may include name of the website: ")
docu = Document()
paragraph = docu.add_paragraph()
length_of_password_requested=int(input("Maximum length of the password is: "))
#for i in range (1,96):
    #print(all_keyboard_entity[i])
all_keyboard_entity={1: '`', 2: '1', 3: '2', 4: '3', 5: '4', 6: '5', 7: '6', 8: '7', 9: '8', 10: 'N', 11: '9', 12: '0', 13: '-', 14: '=', 15: 'q', 16: 'w', 17: 'e', 18: 'r', 19: 't', 20: 'y', 21: 'u', 22: 'i', 23: 'o', 24: 'p', 25: '[', 26: ']', 27: '\\', 28: 'a', 29: 's', 30: 'd', 31: 'f', 32: 'g', 33: 'h', 34: 'j', 35: 'k', 36: 'l', 37: ';', 38: "'", 39: 'z', 40: 'x', 41: 'c', 42: 'v', 43: 'b', 44: 'n', 45: 'm', 46: ',', 47: '.', 48: '/', 49: '~', 50: '!', 51: '@', 52: '#', 53: '$', 54: '%', 55: '^', 56: '&', 57: '*', 58: '', 59: '(', 60: ')', 61: '_', 62: '+', 63: '{', 64: '}', 65: '|', 66: ':', 67: '"', 68: '<', 69: '>', 70: '?', 71: 'Q', 72: 'W', 73: 'E', 74: 'R', 75: 'T', 76: 'Y', 77: 'U', 78: 'I', 79: 'O', 80: 'P', 81: 'A', 82: 'S', 83: 'D', 84: 'F', 85: 'G', 86: 'H', 87: 'J', 88: 'K', 89: 'L', 90: 'Z', 91: 'X', 92: 'C', 93: 'V', 94: 'B', 95: 'M', 96: ' '}
num_of_options=len(all_keyboard_entity)
print("Length of the KeyBoard Array is: ", num_of_options)

PASSWORD=""
for i in range(length_of_password_requested):
    rand=random.randint(1,num_of_options)
    PASSWORD+=all_keyboard_entity[rand]

run = paragraph.add_run('The Password is in within the arrows. Copy the everything within the arrows which is actually in white font. -->')
run.font.color.rgb = RGBColor(255, 0, 0)

run = paragraph.add_run(PASSWORD)
run.font.color.rgb = RGBColor(255,255,255)

run = paragraph.add_run('<-- Dont leave any space while copying everything within the arrows.')
run.font.color.rgb = RGBColor(255, 0, 0)

docu.save('White_Password.docx')
convert('White_Password.docx',intent+"_Password.pdf")

"""
Quantum_Disributor={}
for i in range(1,num_of_options+1):
    Quantum_Disributor[i]=int(0)
# Plot a distribution curve
total=96
for i in range(7):
    total=int(str(total)+str(0))
    for i in range(total):
        rand=random.randint(1,num_of_options)
        Quantum_Disributor.update({rand:Quantum_Disributor[rand]+1})

    plt.figure(figsize=(13,13))
    #plt.figure()
    names = list(Quantum_Disributor.keys())
    values = list(Quantum_Disributor.values())
    S_Dev = np.std(values)
    plt.bar(range(len(Quantum_Disributor)), values, tick_label=names)
    #Mean = np.mean(values)
    #Distance_from_Ideal_Mean=abs((total/96)-Mean)
    #print(Distance_from_Ideal_Mean)

    name="Number of Observations "+str(total)+" with Standard Deviation "+str(S_Dev)
    plt.title(name)
    plt.ylabel("Frequency")
    plt.xlabel("Options_in_PASSWORD")
    add_on=random.randint(0,9999999999)
    name=name+" Unique Image ID "+str(add_on)+".png"
    plt.savefig('Images/{}'.format(name))
    #plt.show()
"""